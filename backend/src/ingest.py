"""Data ingestion pipeline - processes course files into documents."""
import json
from pathlib import Path
from typing import List, Dict
import pandas as pd
from pypdf import PdfReader
from docx import Document as DocxDocument
from tqdm import tqdm

# OCR imports (optional - graceful fallback if not available)
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("[WARN] PyMuPDF not available - OCR support limited. Install with: pip install PyMuPDF")

try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("[WARN] OCR libraries not available - image-based PDFs won't be processed. Install with: pip install pytesseract pdf2image Pillow")

from .config import DATA_RAW_PATH, DOCS_JSONL_PATH
from .models import Document


def extract_text_from_md(file_path: Path) -> str:
    """Extract text from Markdown file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return ""


def extract_text_from_txt(file_path: Path) -> str:
    """Extract text from plain text file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return ""


def extract_text_from_pdf(file_path: Path) -> str:
    """
    Extract text from PDF file with OCR fallback for presentations and scanned PDFs.
    Handles presentation PDFs by marking slide boundaries.
    """
    text = ""
    
    # Strategy 1: Try PyMuPDF (fitz) first - better text extraction and structure detection
    if PYMUPDF_AVAILABLE:
        try:
            doc = fitz.open(str(file_path))
            text_parts = []
            num_pages = len(doc)
            
            empty_slide_indices = []  # Track which slides need OCR
            
            for page_num in range(num_pages):
                page = doc[page_num]
                # Extract text with layout preservation
                page_text = page.get_text("text")
                
                # Check if this slide is image-based (has very little text)
                if len(page_text.strip()) < 20:
                    # This is likely an image-based slide - mark for OCR
                    empty_slide_indices.append(page_num)
                    if page_num == 0:
                        text_parts.append(f"[SLIDE {page_num + 1} - Title/Intro]\n{page_text}\n[IMAGE - OCR Needed]\n")
                    else:
                        text_parts.append(f"[SLIDE {page_num + 1}]\n{page_text}\n[IMAGE - OCR Needed]\n")
                    continue
                
                # Check if this looks like a presentation slide (fewer words, more structure)
                words = page_text.split()
                is_presentation = len(words) < 100  # Presentations typically have less text per slide
                
                if page_num == 0:
                    text_parts.append(f"[SLIDE {page_num + 1} - Title/Intro]\n{page_text}\n")
                elif is_presentation or page_num > 0:
                    text_parts.append(f"[SLIDE {page_num + 1}]\n{page_text}\n")
                else:
                    text_parts.append(f"[PAGE {page_num + 1}]\n{page_text}\n")
            
            text = "\n".join(text_parts)
            
            # If we found empty/image-based slides, try OCR on them
            if empty_slide_indices and OCR_AVAILABLE:
                print(f"[INFO] Found {len(empty_slide_indices)} image-based slide(s): {[i+1 for i in empty_slide_indices]} - attempting OCR...")
                try:
                    from pdf2image import convert_from_path
                    # Convert only the empty slides to images for OCR
                    first_page = min(empty_slide_indices) + 1  # convert_from_path uses 1-based indexing
                    last_page = max(empty_slide_indices) + 1
                    images = convert_from_path(str(file_path), dpi=300, first_page=first_page, last_page=last_page)
                    
                    # Map images back to slide numbers
                    slide_to_image = {empty_slide_indices[idx]: images[idx] for idx in range(len(images)) if idx < len(images)}
                    
                    # Perform OCR on each empty slide
                    for slide_idx in empty_slide_indices:
                        if slide_idx in slide_to_image:
                            slide_num = slide_idx + 1
                            try:
                                ocr_text = pytesseract.image_to_string(slide_to_image[slide_idx], lang='eng')
                                if ocr_text.strip() and len(ocr_text.strip()) > 10:
                                    # Replace the [IMAGE - OCR Needed] marker with OCR text
                                    slide_pattern = f"[SLIDE {slide_num}]\n{page_text}\n[IMAGE - OCR Needed]"
                                    if slide_pattern in text:
                                        text = text.replace(slide_pattern, f"[SLIDE {slide_num}]\n{ocr_text}")
                                    else:
                                        # Try without page_text if it was empty
                                        slide_pattern2 = f"[SLIDE {slide_num}]\n\n[IMAGE - OCR Needed]"
                                        if slide_pattern2 in text:
                                            text = text.replace(slide_pattern2, f"[SLIDE {slide_num}]\n{ocr_text}")
                                        else:
                                            # Find and replace the slide section
                                            import re
                                            pattern = rf'\[SLIDE {slide_num}\][^\[]*\[IMAGE - OCR Needed\]'
                                            replacement = f'[SLIDE {slide_num}]\n{ocr_text}'
                                            text = re.sub(pattern, replacement, text)
                                    print(f"[OCR SUCCESS] Extracted {len(ocr_text)} chars from slide {slide_num}")
                            except Exception as ocr_page_error:
                                print(f"[OCR WARN] Failed to OCR slide {slide_num}: {ocr_page_error}")
                except Exception as ocr_error:
                    print(f"[OCR WARN] OCR processing failed: {ocr_error}")
                    if "tesseract" in str(ocr_error).lower():
                        print("[INFO] Tesseract may not be installed. On Windows, download from: https://github.com/UB-Mannheim/tesseract/wiki")
            
            # Check if we got meaningful text (at least 50 characters per page average)
            
            if text and len(text) > num_pages * 50:
                print(f"[SUCCESS] Extracted {len(text)} chars from PDF using PyMuPDF ({num_pages} pages)")
                if empty_slide_indices:
                    print(f"[INFO] Note: {len(empty_slide_indices)} image-based slide(s) found. OCR attempted.")
                doc.close()
                return text.strip()
            else:
                print(f"[WARN] PyMuPDF extracted minimal text ({len(text)} chars), trying OCR for empty slides...")
                # Don't close yet - will try OCR on empty slides
                if empty_slides and OCR_AVAILABLE:
                    # Try OCR on empty slides
                    try:
                        from pdf2image import convert_from_path
                        images = convert_from_path(str(file_path), dpi=300, first_page=min(empty_slides), last_page=max(empty_slides))
                        for idx, slide_num in enumerate(empty_slides):
                            if idx < len(images):
                                ocr_text = pytesseract.image_to_string(images[idx], lang='eng')
                                if ocr_text.strip():
                                    # Find the slide marker and replace with OCR text
                                    slide_pattern = f"[SLIDE {slide_num}]"
                                    if slide_pattern in text:
                                        text = text.replace(f"{slide_pattern}\n\n", f"{slide_pattern}\n{ocr_text}\n\n")
                                        print(f"[OCR] Extracted {len(ocr_text)} chars from slide {slide_num}")
                    except Exception as ocr_error:
                        print(f"[WARN] OCR on slides failed: {ocr_error}")
                doc.close()
                
        except Exception as e:
            print(f"[WARN] PyMuPDF extraction failed: {e}, trying standard extraction...")
    
    # Strategy 2: Try standard pypdf extraction
    try:
        reader = PdfReader(file_path)
        text_parts = []
        
        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            
            # Mark slide boundaries for presentations
            words = page_text.split()
            is_presentation = len(words) < 100
            
            if page_num == 1:
                text_parts.append(f"[SLIDE {page_num} - Title/Intro]\n{page_text}\n")
            elif is_presentation or page_num > 1:
                text_parts.append(f"[SLIDE {page_num}]\n{page_text}\n")
            else:
                text_parts.append(f"[PAGE {page_num}]\n{page_text}\n")
        
        text = "\n".join(text_parts)
        
        # Check if we got meaningful text
        num_pages = len(reader.pages)
        if text and len(text) > num_pages * 50:
            print(f"[SUCCESS] Extracted {len(text)} chars from PDF using pypdf")
            return text.strip()
        else:
            print(f"[WARN] Standard extraction got minimal text ({len(text)} chars), trying OCR...")
            
    except Exception as e:
        print(f"[WARN] Standard PDF extraction failed: {e}")
    
    # Strategy 3: OCR fallback for scanned/image-based PDFs
    if OCR_AVAILABLE and (not text or len(text) < 100):
        try:
            print(f"[INFO] Attempting OCR extraction for {file_path.name}...")
            
            # Convert PDF pages to images
            images = convert_from_path(str(file_path), dpi=300)
            
            ocr_text_parts = []
            for page_num, image in enumerate(images, 1):
                # Perform OCR on the image
                page_ocr_text = pytesseract.image_to_string(image, lang='eng')
                
                if page_ocr_text.strip():
                    ocr_text_parts.append(f"[SLIDE {page_num} - OCR]\n{page_ocr_text}\n")
                    print(f"[OCR] Extracted {len(page_ocr_text)} chars from page {page_num}")
            
            if ocr_text_parts:
                ocr_text = "\n".join(ocr_text_parts)
                print(f"[SUCCESS] OCR extracted {len(ocr_text)} chars total")
                return ocr_text.strip()
            else:
                print(f"[WARN] OCR returned no text")
                
        except Exception as e:
            print(f"[ERROR] OCR extraction failed: {e}")
            if "tesseract" in str(e).lower():
                print("[INFO] Tesseract may not be installed. On Windows, download from: https://github.com/UB-Mannheim/tesseract/wiki")
    
    # Return whatever we have (could be empty or minimal)
    if not text:
        print(f"[ERROR] Failed to extract any text from PDF: {file_path}")
        return ""
    
    return text.strip()


def extract_text_from_csv(file_path: Path) -> List[str]:
    """Extract text from CSV file - converts each row to a text document."""
    try:
        df = pd.read_csv(file_path)
        texts = []
        for _, row in df.iterrows():
            # Convert row to readable text
            text_parts = []
            for col, val in row.items():
                if pd.notna(val):
                    text_parts.append(f"{col}: {val}")
            if text_parts:
                texts.append("\n".join(text_parts))
        return texts
    except Exception as e:
        print(f"Error reading CSV {file_path}: {e}")
        return []


def format_table_for_extraction(table, table_num: int) -> str:
    """
    Format table into a linear, LLM-friendly format.
    For materials tables, creates a clear list format.
    """
    table_parts = [f"\n[TABLE {table_num}]"]
    
    # Check if this looks like a materials table (has "Class Material", "Required", etc.)
    first_row_text = ""
    if table.rows:
        first_row_cells = [cell.text.strip() for cell in table.rows[0].cells]
        first_row_text = " ".join(first_row_cells).lower()
    
    # Check for grading scale table
    is_grading_table = any(term in first_row_text for term in ["grade", "threshold", "grading scale", "percentage"])
    # Also check table content for grade indicators
    if not is_grading_table and table.rows:
        all_text = " ".join([cell.text.strip() for row in table.rows for cell in row.cells]).lower()
        is_grading_table = any(term in all_text for term in ["a:", "b:", "c:", "d:", "f:", "93", "90", "87", "83", "80"])
    
    is_materials_table = any(term in first_row_text for term in ["class material", "required", "textbook", "course pack"])
    
    if is_grading_table:
        # Special formatting for grading scale table
        table_parts.append("GRADING SCALE:")
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if not row_cells:
                continue
            
            # Format: "Grade: Threshold" or "Grade | Threshold"
            if len(row_cells) >= 2:
                grade = row_cells[0]
                threshold = row_cells[1]
                table_parts.append(f"  {grade}: {threshold}")
            elif len(row_cells) == 1:
                table_parts.append(f"  {row_cells[0]}")
    elif is_materials_table:
        # Special formatting for materials table - extract as clear list
        table_parts.append("REQUIRED CLASS MATERIALS:")
        for row_idx, row in enumerate(table.rows):
            if row_idx == 0:
                continue  # Skip header row
            
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if not row_cells:
                continue
            
            # For materials table, look for "Required" column and material name
            # Format: "Required | Material Name" or just "Material Name"
            if len(row_cells) >= 2:
                # Check if first cell is "Required" or similar
                if row_cells[0].lower() in ["required", "class material"]:
                    material = row_cells[1] if len(row_cells) > 1 else row_cells[0]
                else:
                    # Material name might be in first cell
                    material = row_cells[0]
                
                if material and material.lower() not in ["required", "class material", "class material | class material"]:
                    table_parts.append(f"  • {material}")
            elif len(row_cells) == 1 and row_cells[0]:
                # Single cell with material name
                material = row_cells[0]
                if material.lower() not in ["required", "class material"]:
                    table_parts.append(f"  • {material}")
    else:
        # Standard table formatting
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cell_text = ' '.join(cell_text.split())
                    row_texts.append(cell_text)
            
            if row_texts:
                row_text = " | ".join(row_texts)
                table_parts.append(row_text)
    
    table_parts.append("[END TABLE]\n")
    return "\n".join(table_parts)


def extract_text_from_docx(file_path: Path) -> str:
    """
    Extract text from Word (.docx) file - comprehensive extraction.
    Captures: paragraphs, tables, headers, footers, lists, and images.
    """
    try:
        doc = DocxDocument(file_path)
        text_parts = []
        
        # 1. Extract headers and footers (often contain course info, dates, page numbers)
        for section in doc.sections:
            # Header
            if section.header:
                header_texts = []
                for paragraph in section.header.paragraphs:
                    para_text = paragraph.text.strip()
                    if para_text:
                        header_texts.append(para_text)
                if header_texts:
                    text_parts.append(f"[HEADER] {' | '.join(header_texts)}")
            
            # Footer
            if section.footer:
                footer_texts = []
                for paragraph in section.footer.paragraphs:
                    para_text = paragraph.text.strip()
                    if para_text:
                        footer_texts.append(para_text)
                if footer_texts:
                    text_parts.append(f"[FOOTER] {' | '.join(footer_texts)}")
        
        # 2. Process document body - need to maintain order of paragraphs and tables
        # Use element-based approach to preserve document structure
        body_elements = []
        
        # Track which tables we've processed
        processed_tables = set()
        
        # Iterate through body elements to maintain order
        for element in doc.element.body:
            if element.tag.endswith('p'):
                # It's a paragraph - find matching paragraph object
                for p in doc.paragraphs:
                    if p._element == element:
                        para_text = p.text.strip()
                        if para_text:
                            text_parts.append(para_text)
                            
                            # Check for images in paragraph runs
                            has_image = False
                            for run in p.runs:
                                try:
                                    # Check for drawings/images in run
                                    if hasattr(run._element, 'drawing_lst') and run._element.drawing_lst:
                                        has_image = True
                                    elif hasattr(run._element, 'graphic') and run._element.graphic:
                                        has_image = True
                                except:
                                    pass
                            
                            if has_image:
                                text_parts.append(f"[IMAGE] (in paragraph: {para_text[:50]}...)")
                        break
            
            elif element.tag.endswith('tbl'):
                # It's a table - find matching table object
                for table_idx, table in enumerate(doc.tables):
                    if table._element == element and table._element not in processed_tables:
                        processed_tables.add(table._element)
                        
                        # Extract table with improved formatting
                        table_text = format_table_for_extraction(table, table_idx + 1)
                        if table_text:
                            text_parts.append(table_text)
                        break
        
        # 3. Fallback: Ensure all tables are extracted (in case element matching fails)
        all_table_elements = {t._element for t in doc.tables}
        missed_tables = all_table_elements - processed_tables
        
        if missed_tables:
            for table_idx, table in enumerate(doc.tables):
                if table._element in missed_tables:
                    table_text = format_table_for_extraction(table, table_idx + 1)
                    if table_text:
                        text_parts.append(table_text)
        
        return "\n".join(text_parts)
    
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
        import traceback
        traceback.print_exc()
        return ""


def extract_text_from_file(file_path: Path) -> List[str]:
    """
    Extract text from file based on extension.
    Returns list of text strings (may be multiple for CSV).
    """
    suffix = file_path.suffix.lower()
    
    if suffix == '.md':
        text = extract_text_from_md(file_path)
        return [text] if text else []
    elif suffix == '.txt':
        text = extract_text_from_txt(file_path)
        return [text] if text else []
    elif suffix == '.pdf':
        text = extract_text_from_pdf(file_path)
        return [text] if text else []
    elif suffix == '.docx':
        text = extract_text_from_docx(file_path)
        return [text] if text else []
    elif suffix == '.csv':
        return extract_text_from_csv(file_path)
    else:
        print(f"Unsupported file type: {suffix} for {file_path}")
        return []


def build_breadcrumb(file_path: Path, base_path: Path) -> str:
    """
    Build breadcrumb from file path relative to base.
    Pattern: Modules > Week_X > Filename
    """
    # Get relative path
    rel_path = file_path.relative_to(base_path)
    parts = rel_path.parts
    
    # Remove file extension for display
    filename = file_path.stem
    
    # Build breadcrumb
    if len(parts) > 1:
        # Has subdirectories (e.g., Week_01/file.md)
        dir_parts = parts[:-1]  # All but filename
        breadcrumb = " > ".join(dir_parts)
        return f"{breadcrumb} > {filename}"
    else:
        # Root level file
        return f"Modules > {filename}"


def build_canvas_url(breadcrumb: str) -> str:
    """
    Generate a Canvas URL from breadcrumb (can be customized).
    This is a placeholder - replace with actual Canvas URL mapping.
    """
    # Simple heuristic: convert breadcrumb to URL-friendly path
    # In production, map to actual Canvas URLs
    return f"https://canvas.example.com/modules/{breadcrumb.lower().replace(' > ', '/')}"


def process_directory(base_path: Path = None) -> List[Document]:
    """
    Recursively process all supported files in data/raw directory.
    Returns list of Document objects.
    """
    if base_path is None:
        base_path = DATA_RAW_PATH
    
    documents = []
    supported_extensions = {'.md', '.txt', '.pdf', '.docx', '.csv'}
    
    # Walk through directory
    file_paths = list(base_path.rglob('*'))
    file_paths = [f for f in file_paths if f.is_file() and f.suffix.lower() in supported_extensions]
    
    print(f"Found {len(file_paths)} files to process...")
    
    for file_path in tqdm(file_paths, desc="Processing files"):
        # Extract text(s)
        texts = extract_text_from_file(file_path)
        
        if not texts:
            continue
        
        # Build breadcrumb
        breadcrumb = build_breadcrumb(file_path, base_path)
        
        # Create documents
        for idx, text in enumerate(texts):
            if text.strip():  # Only add non-empty texts
                doc = Document(
                    text=text,
                    breadcrumb=breadcrumb if idx == 0 else f"{breadcrumb} (Row {idx+1})",
                    source_file=str(file_path.relative_to(base_path)),
                    file_type=file_path.suffix.lower(),
                    metadata={
                        "url": build_canvas_url(breadcrumb),
                        "file_path": str(file_path)
                    }
                )
                documents.append(doc)
    
    return documents


def save_to_jsonl(documents: List[Document], output_path: Path = None):
    """Save documents to JSONL format."""
    if output_path is None:
        output_path = DOCS_JSONL_PATH
    
    print(f"Saving {len(documents)} documents to {output_path}...")
    
    with open(output_path, 'w', encoding='utf-8') as f:
        for doc in documents:
            json_line = json.dumps(doc.model_dump(), ensure_ascii=False)
            f.write(json_line + '\n')
    
    print(f"[SUCCESS] Saved {len(documents)} documents to {output_path}")


def main():
    """Main ingestion function."""
    print("=" * 60)
    print("Course Compass - Data Ingestion")
    print("=" * 60)
    
    if not DATA_RAW_PATH.exists():
        print(f"Error: Raw data directory not found: {DATA_RAW_PATH}")
        print(f"Please add course files to {DATA_RAW_PATH}")
        return
    
    # Process files
    documents = process_directory()
    
    if not documents:
        print("No documents found to process!")
        print(f"Please add supported files (.md, .txt, .pdf, .csv) to {DATA_RAW_PATH}")
        return
    
    # Save to JSONL
    save_to_jsonl(documents)
    
    print(f"\n[SUCCESS] Ingestion complete! Processed {len(documents)} documents.")


if __name__ == "__main__":
    main()

